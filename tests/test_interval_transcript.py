from docx import Document

from tw.interval_transcript import fill_interval_table


def _make_table_template(path):
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text, header[3].text = "Time", "Transcription", "Code", "Misc."
    for i in range(1, 4):
        row = table.add_row()
        row.cells[0].text = f"Int {i}"
    doc.save(path)


def test_fill_interval_table_fills_transcription_column_only(tmp_path):
    template = tmp_path / "template.docx"
    _make_table_template(template)
    output = tmp_path / "out.docx"

    buckets = {1: "[00:00:00] [SPEAKER_00] hello", 2: "[00:05:10] [SPEAKER_01] hi back"}
    result_path = fill_interval_table(template, buckets, output)

    assert result_path == output
    doc = Document(output)
    table = doc.tables[0]
    assert table.rows[1].cells[1].text == "[00:00:00] [SPEAKER_00] hello"
    assert table.rows[2].cells[1].text == "[00:05:10] [SPEAKER_01] hi back"


def test_fill_interval_table_removes_rows_with_no_speech(tmp_path):
    template = tmp_path / "template.docx"
    _make_table_template(template)  # Int 1, Int 2, Int 3
    output = tmp_path / "out.docx"

    result_path = fill_interval_table(template, {1: "[00:00:00] [SPEAKER_00] hello"}, output)

    doc = Document(result_path)
    table = doc.tables[0]
    labels = [row.cells[0].text for row in table.rows[1:]]
    assert labels == ["Int 1"]
    assert table.rows[1].cells[1].text == "[00:00:00] [SPEAKER_00] hello"


def test_fill_interval_table_appends_rows_beyond_template(tmp_path):
    template = tmp_path / "template.docx"
    _make_table_template(template)  # Int 1, Int 2, Int 3
    output = tmp_path / "out.docx"

    buckets = {
        1: "[00:00:00] [SPEAKER_00] hello",
        2: "[00:05:00] [SPEAKER_00] hi",
        3: "[00:10:00] [SPEAKER_00] hey",
        4: "[00:15:00] [SPEAKER_00] extra one",
        5: "[00:20:00] [SPEAKER_00] extra two",
    }
    result_path = fill_interval_table(template, buckets, output)

    doc = Document(result_path)
    table = doc.tables[0]
    labels = [row.cells[0].text for row in table.rows[1:]]
    assert labels == ["Int 1", "Int 2", "Int 3", "Int 4", "Int 5"]
    assert table.rows[4].cells[1].text == "[00:15:00] [SPEAKER_00] extra one"
    assert table.rows[5].cells[1].text == "[00:20:00] [SPEAKER_00] extra two"


def test_fill_interval_table_leaves_other_columns_untouched(tmp_path):
    template = tmp_path / "template.docx"
    _make_table_template(template)
    table = Document(template).tables[0]
    doc = Document(template)
    doc.tables[0].rows[1].cells[2].text = "existing code note"
    doc.save(template)
    output = tmp_path / "out.docx"

    fill_interval_table(template, {1: "hello"}, output)

    doc = Document(output)
    assert doc.tables[0].rows[1].cells[2].text == "existing code note"
    assert doc.tables[0].rows[1].cells[0].text == "Int 1"
