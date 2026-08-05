from docx import Document

from tw.speaker_rename import extract_speaker_labels, extract_speakers_from_file, rename_speakers_in_file, rename_speakers_in_text


def test_extract_speaker_labels_finds_distinct_tags():
    text = "[SPEAKER_00] hello\n\n[SPEAKER_01] hi\n\n[SPEAKER_00] again\n\n[UNKNOWN] mystery"
    assert extract_speaker_labels(text) == {"SPEAKER_00", "SPEAKER_01", "UNKNOWN"}


def test_rename_speakers_in_text_replaces_tags():
    text = "[SPEAKER_00] hello\n\n[SPEAKER_01] hi back"
    result = rename_speakers_in_text(text, {"SPEAKER_00": "Mom", "SPEAKER_01": "Dad"})
    assert result == "[Mom] hello\n\n[Dad] hi back"


def test_rename_speakers_in_text_leaves_unmapped_labels_alone():
    text = "[SPEAKER_00] hello\n\n[SPEAKER_01] hi"
    result = rename_speakers_in_text(text, {"SPEAKER_00": "Mom"})
    assert result == "[Mom] hello\n\n[SPEAKER_01] hi"


def test_extract_speakers_from_markdown_file(tmp_path):
    md_path = tmp_path / "out.md"
    md_path.write_text("## Raw Transcript\n\n[SPEAKER_00] hi\n\n[SPEAKER_01] hey")
    assert extract_speakers_from_file(md_path) == {"SPEAKER_00", "SPEAKER_01"}


def test_rename_speakers_in_markdown_file(tmp_path):
    md_path = tmp_path / "out.md"
    md_path.write_text("[SPEAKER_00] hi")
    output = tmp_path / "renamed.md"

    rename_speakers_in_file(md_path, {"SPEAKER_00": "Mom"}, output)

    assert output.read_text() == "[Mom] hi"


def test_extract_speakers_from_docx_table(tmp_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "[SPEAKER_00] hi\n\n[SPEAKER_01] hey"
    docx_path = tmp_path / "out.docx"
    doc.save(docx_path)

    assert extract_speakers_from_file(docx_path) == {"SPEAKER_00", "SPEAKER_01"}


def test_rename_speakers_in_docx_table_cell(tmp_path):
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Int 1"
    table.rows[0].cells[1].text = "[SPEAKER_00] hi"
    docx_path = tmp_path / "out.docx"
    doc.save(docx_path)
    output = tmp_path / "renamed.docx"

    rename_speakers_in_file(docx_path, {"SPEAKER_00": "Mom"}, output)

    result = Document(output)
    assert result.tables[0].rows[0].cells[1].text == "[Mom] hi"
    assert result.tables[0].rows[0].cells[0].text == "Int 1"


def test_rename_speakers_in_docx_paragraph(tmp_path):
    doc = Document()
    doc.add_paragraph("[SPEAKER_00] hi there")
    docx_path = tmp_path / "out.docx"
    doc.save(docx_path)
    output = tmp_path / "renamed.docx"

    rename_speakers_in_file(docx_path, {"SPEAKER_00": "Mom"}, output)

    result = Document(output)
    assert result.paragraphs[0].text == "[Mom] hi there"
