from docx import Document

from tw.docx_template import parse_fields_docx, assemble_docx


def _make_template(path, headings):
    doc = Document()
    for level, text in headings:
        doc.add_heading(text, level=level)
    doc.save(path)


def test_parse_fields_docx_reads_heading_paragraphs(tmp_path):
    template = tmp_path / "template.docx"
    _make_template(template, [(1, "Client Call Notes"), (2, "Attendees"), (2, "Decisions")])
    assert parse_fields_docx(template) == ["Client Call Notes", "Attendees", "Decisions"]


def test_parse_fields_docx_ignores_non_heading_paragraphs(tmp_path):
    doc = Document()
    doc.add_heading("Notes", level=2)
    doc.add_paragraph("Just a normal paragraph, not a heading.")
    template = tmp_path / "template.docx"
    doc.save(template)
    assert parse_fields_docx(template) == ["Notes"]


def test_assemble_docx_fills_all_fields(tmp_path):
    template = tmp_path / "template.docx"
    _make_template(template, [(2, "Attendees"), (2, "Decisions")])
    output = tmp_path / "out.docx"

    result_path = assemble_docx(template, {"Attendees": "Alice, Bob", "Decisions": "Ship v1"}, [], output)

    assert result_path == output
    doc = Document(output)
    paragraphs = [p.text for p in doc.paragraphs]
    assert "Attendees" in paragraphs
    assert "Alice, Bob" in paragraphs
    assert "Decisions" in paragraphs
    assert "Ship v1" in paragraphs


def test_assemble_docx_marks_failed_fields(tmp_path):
    template = tmp_path / "template.docx"
    _make_template(template, [(2, "Attendees"), (2, "Decisions")])
    output = tmp_path / "out.docx"

    assemble_docx(template, {"Attendees": "Alice, Bob"}, ["Decisions"], output)

    doc = Document(output)
    paragraphs = [p.text for p in doc.paragraphs]
    assert "[NEEDS MANUAL REVIEW]" in paragraphs
    assert any("NEEDS MANUAL REVIEW" in p and "Decisions" in p for p in paragraphs[:1])


def test_assemble_docx_raises_on_unaccounted_field(tmp_path):
    import pytest

    template = tmp_path / "template.docx"
    _make_template(template, [(2, "Attendees"), (2, "Decisions")])
    output = tmp_path / "out.docx"

    with pytest.raises(ValueError, match="Decisions"):
        assemble_docx(template, {"Attendees": "Alice"}, [], output)
