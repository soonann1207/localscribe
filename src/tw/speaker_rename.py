import re
from pathlib import Path

from docx import Document

_SPEAKER_TAG_RE = re.compile(r"\[(SPEAKER_\d+|UNKNOWN)\]")


def extract_speaker_labels(text: str) -> set[str]:
    return set(_SPEAKER_TAG_RE.findall(text))


def rename_speakers_in_text(text: str, mapping: dict[str, str]) -> str:
    return _SPEAKER_TAG_RE.sub(lambda m: f"[{mapping.get(m.group(1), m.group(1))}]", text)


def _docx_all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_speakers_from_file(path: Path) -> set[str]:
    if path.suffix.lower() == ".docx":
        return extract_speaker_labels(_docx_all_text(Document(path)))
    return extract_speaker_labels(path.read_text())


def rename_speakers_in_file(path: Path, mapping: dict[str, str], output_path: Path) -> Path:
    if path.suffix.lower() == ".docx":
        doc = Document(path)
        for p in doc.paragraphs:
            for run in p.runs:
                run.text = rename_speakers_in_text(run.text, mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = rename_speakers_in_text(cell.text, mapping)
        doc.save(output_path)
        return output_path

    output_path.write_text(rename_speakers_in_text(path.read_text(), mapping))
    return output_path
