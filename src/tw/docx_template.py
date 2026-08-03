from pathlib import Path

from docx import Document


def parse_fields_docx(template_path: Path) -> list[str]:
    doc = Document(template_path)
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


def _heading_level(style_name: str) -> int:
    return int(style_name.rsplit(" ", 1)[-1])


def assemble_docx(
    template_path: Path,
    filled: dict[str, str],
    failed_fields: list[str],
    output_path: Path,
) -> Path:
    src = Document(template_path)
    headings = [
        (p.text, _heading_level(p.style.name))
        for p in src.paragraphs
        if p.style.name.startswith("Heading")
    ]

    out = Document()
    if failed_fields:
        out.add_paragraph("NEEDS MANUAL REVIEW: " + ", ".join(failed_fields))

    for field_name, level in headings:
        if field_name in failed_fields:
            content = "[NEEDS MANUAL REVIEW]"
        elif field_name in filled:
            content = filled[field_name]
        else:
            raise ValueError(f"template field '{field_name}' not in filled or failed_fields")
        out.add_heading(field_name, level=level)
        out.add_paragraph(content)

    out.save(output_path)
    return output_path
